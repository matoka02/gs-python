import os
import shutil
import pandas as pd
from fastapi import Depends, APIRouter, Form, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfReader, PdfWriter

from .db import User
from .schemas import ImageProcessingOptions
from .users import current_active_user
from .utils import process_image

file_router = APIRouter()
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# --- IMAGES ---

@file_router.post("/upload/images")
async def upload_file(
        file: UploadFile = File(...),
        resize: str = Form(None),
        convert_to: str = Form(None),
        grayscale: bool = Form(None),
        flip: str = Form(None),
        user: User = Depends(current_active_user)
):
    options = ImageProcessingOptions(
        resize=resize,
        convert_to=convert_to,
        grayscale=grayscale,
        flip=flip
    )
    print(options)
    # создание новой папки для хранения
    user_folder = f"{UPLOAD_DIR}/{user.id}"
    os.makedirs(user_folder, exist_ok=True)
    # указание папки хранения
    file_location = os.path.join(user_folder, file.filename)
    # открыть файл
    with open(file_location, 'wb') as buffers:
        shutil.copyfileobj(file.file, buffers)
    # функция обработки изображений
    processed_file_location = process_image(file_location, options)
    return {'filename': os.path.basename(processed_file_location), 'location': processed_file_location}


@file_router.get("/download/{filename}", response_class=FileResponse)
async def download_file(
        filename: str,
        user: User = Depends(current_active_user)
):
    user_folder = f"{UPLOAD_DIR}/{user.id}"
    # указание папки хранения
    file_location = os.path.join(user_folder, filename)
    if not os.path.exists(file_location):
        raise HTTPException(status_code=404, detail='File not found')
    # путь загрузки
    return FileResponse(file_location)


# --- CSV to XLSX and XLSX to CSV ---

@file_router.post("/convert/csv-to-xlsx/")
async def convert_csv_to_xlsx(file: UploadFile = File(...), user: User = Depends(current_active_user)):
    user_folder = f"{UPLOAD_DIR}/{user.id}"
    os.makedirs(user_folder, exist_ok=True)

    csv_file_path = os.path.join(user_folder, file.filename)
    with open(csv_file_path, 'wb') as f:
        shutil.copyfileobj(file.file, f)

    try:
        # конвертация CSV в XLSX
        df = pd.read_csv(csv_file_path)
        xlsx_file_path = os.path.splitext(csv_file_path)[0] + ".xlsx"
        df.to_excel(xlsx_file_path, index=False)
        # удаление исходного CSV-файла
        os.remove(csv_file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting CSV to XLSX: {str(e)}")

    return {"filename": os.path.basename(xlsx_file_path), "location": xlsx_file_path}


@file_router.post("/convert/xlsx-to-csv/")
async def convert_xlsx_to_csv(file: UploadFile = File(...), user: User = Depends(current_active_user)):
    user_folder = f"{UPLOAD_DIR}/{user.id}"
    os.makedirs(user_folder, exist_ok=True)

    xlsx_file_path = os.path.join(user_folder, file.filename)
    with open(xlsx_file_path, 'wb') as f:
        shutil.copyfileobj(file.file, f)

    try:
        # конвертация XLSX в CSV
        df = pd.read_excel(xlsx_file_path)
        csv_file_path = os.path.splitext(xlsx_file_path)[0] + ".csv"
        df.to_csv(csv_file_path, index=False)
        # удаление исходного XLSX-файла
        os.remove(xlsx_file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting XLSX to CSV: {str(e)}")

    return {"filename": os.path.basename(csv_file_path), "location": csv_file_path}


# --- DOCX/DOC to PDF and PDF to DOCX ---

@file_router.post("/convert/doc-to-pdf/")
async def convert_doc_to_pdf(file: UploadFile = File(...), user: User = Depends(current_active_user)):
    user_folder = f"{UPLOAD_DIR}/{user.id}"
    os.makedirs(user_folder, exist_ok=True)

    doc_file_path = os.path.join(user_folder, file.filename)
    with open(doc_file_path, 'wb') as f:
        shutil.copyfileobj(file.file, f)

    try:
        pdf_file_path = os.path.splitext(doc_file_path)[0] + ".pdf"
        doc = Document(doc_file_path)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        pdf.output(pdf_file_path)
        # удаление исходного DOCX-файла
        os.remove(doc_file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting DOC/DOCX to PDF: {str(e)}")

    return {"filename": os.path.basename(pdf_file_path), "location": pdf_file_path}


@file_router.post("/convert/pdf-to-doc/")
async def convert_pdf_to_doc(file: UploadFile = File(...), user: User = Depends(current_active_user)):
    user_folder = f"{UPLOAD_DIR}/{user.id}"
    os.makedirs(user_folder, exist_ok=True)

    pdf_file_path = os.path.join(user_folder, file.filename)
    with open(pdf_file_path, 'wb') as f:
        shutil.copyfileobj(file.file, f)

    try:
        doc_file_path = os.path.splitext(pdf_file_path)[0] + ".docx"
        reader = PdfReader(pdf_file_path)
        doc = Document()
        for page in reader.pages:
            doc.add_paragraph(page.extract_text())
        doc.save(doc_file_path)
        # удаление исходного PDF-файла
        os.remove(pdf_file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting PDF to DOC/DOCX: {str(e)}")

    return {"filename": os.path.basename(doc_file_path), "location": doc_file_path}